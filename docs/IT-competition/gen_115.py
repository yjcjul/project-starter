import os
from docx import Document
import openpyxl
from openpyxl import load_workbook
import copy

SRC_DIR = r'C:\projects\project-starter\docs\IT-competition\references'
OUT_DIR = r'C:\projects\project-starter\docs\IT-competition'

students = [
    {'班級': '807', '座號': '10', '姓名': '黃云毅', '項目': '電腦打字國中英文甲組',          '獎等': '優等', '嘉獎': '小功壹次', '指導老師': '李倩瑜'},
    {'班級': '709', '座號': '16', '姓名': '蕭睿彭', '項目': 'SCRATCH應用競賽國中甲組生活應用', '獎等': '甲等', '嘉獎': '嘉獎貳次', '指導老師': '陳怡蓁'},
    {'班級': '701', '座號': '3',  '姓名': '呂善衡', '項目': 'SCRATCH應用競賽國中甲組生活應用', '獎等': '甲等', '嘉獎': '嘉獎貳次', '指導老師': '陳怡蓁'},
]

prize = {'黃云毅': 200, '蕭睿彭': 150, '呂善衡': 150}

# ── 1. 獎懲申請表（每人一份）─────────────────────────────
src_docx = os.path.join(SRC_DIR, '獎懲申請表.docx')

def replace_in_doc(doc, replacements):
    for p in doc.paragraphs:
        for old, new in replacements:
            if old in p.text:
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements:
                        if old in p.text:
                            for run in p.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

for s in students:
    d = Document(src_docx)
    replacements = [
        ('114年', '115年'),
        ('05月13日', '  月  日'),
        ('707', s['班級']),
        ('黃云毅', s['姓名']),
        ('電腦打字國中甲組', s['項目']),
        ('優等', s['獎等']),
        ('小功壹次', s['嘉獎']),
    ]
    # 座號：只換確定是座號的格子（值為 "10"）
    replace_in_doc(d, replacements)
    # 座號欄位單獨處理
    for t in d.tables:
        for row in t.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                if cell.text.strip() == '10':
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = run.text.replace('10', s['座號'])
    out = os.path.join(OUT_DIR, f'115年_獎懲申請表_{s["姓名"]}.docx')
    d.save(out)
    print(f'✅ {out}')

# ── 2. 敘獎擬議清冊（Excel）─────────────────────────────
src_xlsx = os.path.join(SRC_DIR, '敘獎擬議清冊.xlsx')
wb = load_workbook(src_xlsx)
ws = wb.active

teachers = [
    ('教師', '李倩瑜', '指導學生獲電腦打字國中英文甲組優等', '嘉獎1次'),
    ('教師', '陳怡蓁', '指導學生獲SCRATCH應用競賽國中甲組生活應用甲等', '嘉獎1次'),
]

for row in ws.iter_rows():
    for cell in row:
        if cell.value and '114年度' in str(cell.value):
            cell.value = cell.value.replace('114年度', '115年度')
        if cell.value and '1140043071' in str(cell.value):
            cell.value = '來文依據：（待填入115年公文字號）'
        if cell.value and '規定可敘獎總人數' in str(cell.value):
            cell.value = f'規定可敘獎總人數： {len(teachers)} 人'

# 清除舊教師資料並填入新資料
data_start = None
for i, row in enumerate(ws.iter_rows(), 1):
    if row[0].value == '職稱':
        data_start = i + 1
        break

if data_start:
    for j, (title, name, work, award) in enumerate(teachers):
        r = data_start + j
        ws.cell(row=r, column=1, value=title)
        ws.cell(row=r, column=2, value=name)
        ws.cell(row=r, column=3, value=work)
        ws.cell(row=r, column=4, value=award)

out_xlsx = os.path.join(OUT_DIR, '115年_敘獎擬議清冊.xlsx')
wb.save(out_xlsx)
print(f'✅ {out_xlsx}')

# ── 3. 家長會獎學金印領清冊（Word）──────────────────────
src_money = os.path.join(SRC_DIR, '家長會獎學金印領清冊-114年資訊競賽.docx')
d = Document(src_money)

replace_in_doc(d, [('114年度', '115年度'), ('114年', '115年')])

# 找到資料表格並清除舊資料、填入新資料
for t in d.tables:
    # 找有「編號申請項目」的表格
    header_found = any('申請項目' in cell.text for row in t.rows for cell in row.cells)
    if not header_found:
        continue
    # 找資料起始列
    data_rows = []
    for i, row in enumerate(t.rows):
        if any('申請項目' in cell.text or '名次' in cell.text for cell in row.cells):
            data_start_row = i + 1
            break
    # 清除現有資料列
    for row in t.rows[data_start_row:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ''

    # 重新填入 115 資料
    # 學生
    rows_needed = [
        ('115年度中小學資訊網路應用競賽-「電腦打字國中英文甲組」',    '優等', '807', '黃云毅',   '200元'),
        ('115年度中小學資訊網路應用競賽-「電腦打字國中英文甲組」',    '優等', '指導教師', '李倩瑜', '200元'),
        ('115年度中小學資訊網路應用競賽-「SCRATCH應用競賽國中甲組生活應用」', '甲等', '709', '蕭睿彭',   '150元'),
        ('115年度中小學資訊網路應用競賽-「SCRATCH應用競賽國中甲組生活應用」', '甲等', '指導教師', '陳怡蓁', '150元'),
        ('115年度中小學資訊網路應用競賽-「SCRATCH應用競賽國中甲組生活應用」', '甲等', '701', '呂善衡',   '150元'),
    ]

    for i, (item, rank, cls, name, money) in enumerate(rows_needed):
        r_idx = data_start_row + i
        if r_idx < len(t.rows):
            row = t.rows[r_idx]
            vals = [item, rank, cls, name, money, '']
            for j, val in enumerate(vals):
                if j < len(row.cells):
                    cell = row.cells[j]
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = ''
                    if cell.paragraphs:
                        cell.paragraphs[0].add_run(val)

    # 更新總計
    for row in t.rows:
        for cell in row.cells:
            if '700元整' in cell.text:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.text = run.text.replace('700元整', '850元整')

out_money = os.path.join(OUT_DIR, '115年_家長會獎學金印領清冊.docx')
d.save(out_money)
print(f'✅ {out_money}')

print('\n🎉 全部文件生成完成！')
